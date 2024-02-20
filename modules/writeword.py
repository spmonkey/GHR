'''
Function:
    将结果写入word文档
Author:
    花果山
Wechat official account：
    中龙 红客突击队
Official website：
    https://www.hscsec.cn/
Email：
    spmonkey@hscsec.cn
Blog:
    https://spmonkey.github.io/
GitHub:
    https://github.com/spmonkey/
'''
# -*- coding: utf-8 -*-
import docx
import os
import re
import sys
from urllib.parse import urlparse
from docx.oxml.ns import qn


class WW:
    def __init__(self, url, result_list):
        self.file_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.url = url
        filename = urlparse(self.url)
        netloc = "_".join(filename.netloc.split("."))
        scheme = filename.scheme
        path = filename.path
        if ":" in netloc:
            netloc = netloc.replace(":", "_")
        if "/" in path:
            path = path.replace("/", "_")
        self.file_name = "{}_{}{}".format(scheme, netloc, path)
        self.result = result_list
        self.result_list = []
        self.vuln_path = []
        self.vuln_request = []
        self.vuln_descriptions = {}
        self.result_dicts = {}
        self.repair_suggestions = {}

    def sort(self):
        for result in self.result:
            if "OPTIONS" in result:
                result = result.replace("开启了 OPTIONS 方法", "检测到目标站点存在开启了 OPTIONS 方法漏洞")
            vuln_name = re.findall("存在(.*)漏洞(.*)\x1b", result)
            if "CVE" in vuln_name[0][1]:
                vulnname = vuln_name[0][0] + ":" + vuln_name[0][1]
            else:
                vulnname = vuln_name[0][0]
            try:
                vuln_path = re.search("(/.*)\sHTTP/1.1", result).group(1)
            except:
                vuln_path = "无"
            vuln_request = re.sub("(.*)漏洞(.*)\x1b\[0m\n", "", result)
            vuln_request = re.sub("                 ", "", vuln_request)
            result_dir = {vulnname: {"漏洞位置": vuln_path, "漏洞请求": vuln_request}}
            self.result_list.append(result_dir)
        return True

    def neaten(self):
        if sys.platform.startswith("win"):
            vuln_descriptions = open("{}\\library\\vuln_description.txt".format(self.file_path), "r", encoding="utf-8").readlines()
        else:
            vuln_descriptions = open("{}/library/vuln_description.txt".format(self.file_path), "r", encoding="utf-8").readlines()
        for vuln_description in vuln_descriptions:
            all = vuln_description.strip().split("：")
            self.vuln_descriptions[all[0]] = all[1]
        for result_list in self.result_list:
            for vulnname in result_list.keys():
                if vulnname not in self.result_dicts.keys():
                    self.result_dicts[vulnname] = {"漏洞位置": [result_list[vulnname]["漏洞位置"]], "漏洞请求": [result_list[vulnname]["漏洞请求"]]}
                else:
                    self.result_dicts[vulnname]["漏洞位置"].append(result_list[vulnname]["漏洞位置"])
                    self.result_dicts[vulnname]["漏洞请求"].append(result_list[vulnname]["漏洞请求"])
        if sys.platform.startswith("win"):
            repair_suggestions = open("{}\\library\\repair_suggestion.txt".format(self.file_path), "r", encoding="utf-8").readlines()
        else:
            repair_suggestions = open("{}/library/repair_suggestion.txt".format(self.file_path), "r", encoding="utf-8").readlines()
        for repair_suggestion in repair_suggestions:
            all = repair_suggestion.strip().split(":")
            self.repair_suggestions[all[0]] = all[1]
        return True

    def writeword(self):
        if sys.platform.startswith("win"):
            full_path = os.path.join(self.file_path + "\\report", self.file_name + ".docx")
        else:
            full_path = os.path.join(self.file_path + "/report", self.file_name + ".docx")
        if os.path.exists(full_path):
            os.remove(full_path)
        doc = docx.Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        url = doc.add_heading("", level=1)
        url_run = url.add_run(self.url)
        url_run.font.name = u'宋体'
        for vulnname in self.result_dicts.keys():
            vuln_name_add = doc.add_heading("", level=2)
            if "CVE" in vulnname:
                vuln = vulnname.split(":")[0]
                cve = vulnname.split(":")[1]
            else:
                vuln = vulnname
            if "OPTIONS" in vuln:
                vuln_name_run = vuln_name_add.add_run(vuln)
            else:
                if ":" in vulnname:
                    vuln_name_run = vuln_name_add.add_run(vuln + "漏洞" + cve)
                else:
                    vuln_name_run = vuln_name_add.add_run(vuln + "漏洞")
            vuln_name_run.font.name = u'宋体'
            vuln_name_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            describe = doc.add_heading("", level=3)
            describe_run = describe.add_run("漏洞描述")
            describe_run.font.name = u'宋体'
            describe_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            if "OPTIONS" in vuln:
                doc.add_paragraph(self.vuln_descriptions[vuln])
            else:
                doc.add_paragraph(self.vuln_descriptions[vuln+"漏洞"])
            path = doc.add_heading("", level=3)
            path_run = path.add_run("漏洞位置")
            path_run.font.name = u'宋体'
            path_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            for vuln_path in self.result_dicts[vulnname]["漏洞位置"]:
                doc.add_paragraph("{}".format(vuln_path))
            req = doc.add_heading("", level=3)
            req_run = req.add_run("漏洞请求")
            req_run.font.name = u'宋体'
            req_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            for vuln_request in self.result_dicts[vulnname]["漏洞请求"]:
                doc.add_paragraph(vuln_request)
            if sys.platform.startswith("win"):
                doc.save(self.file_path+"\\report\\{}.docx".format(self.file_name))
            else:
                doc.save(self.file_path+"/report/{}.docx".format(self.file_name))
            sug = doc.add_heading("", level=3)
            sug_run = sug.add_run("修复建议")
            sug_run.font.name = u'宋体'
            sug_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            if "OPTIONS" in vuln:
                for repair_suggestion in self.repair_suggestions[vuln].split("\\n"):
                    doc.add_paragraph(repair_suggestion)
            else:
                for repair_suggestion in self.repair_suggestions[vuln + "漏洞"].split("\\n"):
                    doc.add_paragraph(repair_suggestion)
            if sys.platform.startswith("win"):
                doc.save(self.file_path+"\\report\\{}.docx".format(self.file_name))
            else:
                doc.save(self.file_path+"/report/{}.docx".format(self.file_name))
        if sys.platform.startswith("win"):
            print("报告生成成功，报告路径为：{}".format(self.file_path+"\\report\\{}.docx".format(self.file_name)))
        else:
            print("报告生成成功，报告路径为：{}".format(self.file_path+"/report/{}.docx".format(self.file_name)))

    def main(self):
        if self.sort():
            if self.neaten():
                self.writeword()

